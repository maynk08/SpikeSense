"""
Finalize the report: build it, render to PDF, compute the real page number of
every heading, figure and table, then inject a populated Table of Contents,
List of Figures and List of Tables (dot leaders + page numbers), and re-render.

Static lists are generated rather than relying on Word fields, so they display
correctly in any viewer and in the submitted PDF. Two passes keep page numbers
correct: the lists are inserted first with blank numbers (fixing the final
length), the PDF is rendered to read true pages, then the same entries are
rewritten with real numbers — leaving pagination unchanged.

Run:  python report/scripts/finalize_report.py
"""

from __future__ import annotations

import json
import re
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

import build_report as B

ROOT = B.ROOT
DOCX = B.OUT
PDF = ROOT / "report" / "Spike-Sense_Project_Report.pdf"
INDEX = ROOT / "report" / "_index.json"
SOFFICE = "/Applications/LibreOffice.app/Contents/MacOS/soffice"
PDFTOTEXT = "/opt/homebrew/bin/pdftotext"
TOC_RIGHT = 6.3  # inches


def render_pdf():
    subprocess.run(
        [SOFFICE, "--headless", "--convert-to", "pdf", "--outdir", str(PDF.parent), str(DOCX)],
        check=True, capture_output=True,
    )


def page_texts():
    out = subprocess.run([PDFTOTEXT, "-layout", str(PDF), "-"],
                         check=True, capture_output=True, text=True).stdout
    pages = []
    for raw in out.split("\f"):
        m = re.search(r"Page (\d+) of \d+", raw)
        pages.append((int(m.group(1)) if m else None, re.sub(r"[ \t]+", " ", raw)))
    return pages


def collect_toc_entries(doc):
    entries = []
    for p in doc.paragraphs:
        s, t = p.style.name, p.text.strip()
        if not t or t == "Table of Contents":
            continue
        if s == "Heading 1":
            entries.append((t, 1))
        elif s == "Heading 2":
            entries.append((t, 2))
    return entries


def heading_pages(toc_entries):
    """Map heading title -> footer page, skipping the TOC pages themselves.
    Headings are searched in document order with a non-decreasing page cursor, so
    a short title (e.g. 'Abstract') is not matched against an earlier, longer
    heading that merely contains it (e.g. 'Abstract / Executive Summary')."""
    pages = page_texts()
    titles = [t for t, _ in toc_entries]
    toc_pages = {i for i, (_, txt) in enumerate(pages)
                 if sum(1 for t in titles if t in re.sub(r"\.{2,}\s*\d+", "", txt)) >= 5}
    out = {}
    cursor = 0
    for title, _ in toc_entries:
        for i in range(cursor, len(pages)):
            if i in toc_pages:
                continue
            footer, txt = pages[i]
            if title in txt:
                out[title] = footer if footer is not None else ""
                cursor = i
                break
    return out


def label_pages(labels):
    """Map a figure/table label -> footer page by searching for 'Label:' (the
    colon distinguishes the real caption from the list-of-figures entry)."""
    pages = page_texts()
    out = {}
    for label in labels:
        needle = label + ":"
        for footer, txt in pages:
            if needle in txt:
                out[label] = footer if footer is not None else ""
                break
    return out


def _entry_paragraph(doc, text, page, indent=0.0, bold=False):
    para = doc.add_paragraph()
    pf = para.paragraph_format
    pf.line_spacing = 1.3
    pf.space_after = Pt(2)
    if indent:
        pf.left_indent = Inches(indent)
    pf.tab_stops.add_tab_stop(Inches(TOC_RIGHT), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    r1 = para.add_run(text)
    B.set_run_font(r1, size=12, bold=bold)
    r2 = para.add_run("\t" + str(page))
    B.set_run_font(r2, size=12, bold=bold)
    return para


def replace_marker(doc, marker, entries):
    """Replace the marker paragraph with the given entry paragraphs.
    entries: list of (text, page, indent, bold)."""
    anchor = None
    for p in doc.paragraphs:
        if p.text.strip() == marker:
            anchor = p
            break
    if anchor is None:
        return
    for text, page, indent, bold in entries:
        para = _entry_paragraph(doc, text, page, indent, bold)
        anchor._p.addprevious(para._p)
    anchor._p.getparent().remove(anchor._p)


def fill_lists(toc_entries, fig_index, tbl_index, hpages, fpages, tpages):
    doc = Document(str(DOCX))
    replace_marker(doc, B.TOC_MARKER, [
        (title, hpages.get(title, ""), 0.0 if lvl == 1 else 0.3, lvl == 1)
        for title, lvl in toc_entries
    ])
    replace_marker(doc, B.LOF_MARKER, [
        (f"{label}  {caption.rstrip('.')}", fpages.get(label, ""), 0.0, False)
        for label, caption in fig_index
    ])
    replace_marker(doc, B.LOT_MARKER, [
        (f"{label}  {caption.rstrip('.')}", tpages.get(label, ""), 0.0, False)
        for label, caption in tbl_index
    ])
    B.force_font(doc)
    doc.save(str(DOCX))


def main():
    idx = json.loads(INDEX.read_text()) if INDEX.exists() else {"figures": [], "tables": []}
    fig_index = idx["figures"]
    tbl_index = idx["tables"]
    fig_labels = [l for l, _ in fig_index]
    tbl_labels = [l for l, _ in tbl_index]

    # Pass 1: build, fill lists with blank pages so the length is final.
    B.main()
    toc_entries = collect_toc_entries(Document(str(DOCX)))
    blank = {t: "" for t, _ in toc_entries}
    fill_lists(toc_entries, fig_index, tbl_index, blank,
               {l: "" for l in fig_labels}, {l: "" for l in tbl_labels})

    # Render and read true pages.
    render_pdf()
    hpages = heading_pages(toc_entries)
    fpages = label_pages(fig_labels)
    tpages = label_pages(tbl_labels)

    # Pass 2: rebuild and fill with real page numbers (identical line counts).
    B.main()
    fill_lists(toc_entries, fig_index, tbl_index, hpages, fpages, tpages)
    render_pdf()

    print(f"Finalized {DOCX.name} + {PDF.name}")
    print(f"  TOC:{len(toc_entries)} mapped:{len(hpages)} | "
          f"figures:{len(fig_index)} mapped:{len(fpages)} | "
          f"tables:{len(tbl_index)} mapped:{len(tpages)}")


if __name__ == "__main__":
    main()
