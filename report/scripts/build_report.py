"""
Build the Spike-Sense MCA project report into the university .docx template.

Strategy
--------
The template ships with placeholder paragraphs ("Write content here...") under
each chapter heading, a cover page inside a content control (w:sdt), and a
Declaration with blanks. This script:

  1. Fills the cover-page fields and the Declaration / Acknowledgement.
  2. Replaces each "Write content here..." placeholder (Abstract + 10 chapters)
     with rich content (headings, paragraphs, bullet lists, numbered figures,
     tables and code listings) defined in content.py.
  3. Applies the required formatting: Times New Roman 12 pt body, 1.5 line
     spacing, justified, with numbered/captioned figures and tables.

All prose is original; figures are generated from the real project artifacts.

Run:  python report/scripts/build_report.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor

import content as C

ROOT = Path(__file__).resolve().parents[2]
TEMPLATE = Path("/Users/mayank/Downloads/Project Report Template.docx")
FIGDIR = ROOT / "report" / "figures"
OUT = ROOT / "report" / "Spike-Sense_Project_Report.docx"

BODY_FONT = "Times New Roman"
BODY_SIZE = 12
MONO_FONT = "Consolas"


# ---------------------------------------------------------------------------
# Low-level helpers
# ---------------------------------------------------------------------------

def set_run_font(run, name=BODY_FONT, size=BODY_SIZE, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    # ensure east-asian / hAnsi also use the font
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(attr), name)


def style_paragraph(p, size=BODY_SIZE, justify=True, spacing=1.5, space_after=8):
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = spacing
    pf.space_after = Pt(space_after)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def set_element_text(p_el, text):
    """Set a w:p element's text via its first run, clearing extra runs."""
    runs = p_el.findall(qn("w:r"))
    if not runs:
        r = OxmlElement("w:r")
        p_el.append(r)
        runs = [r]
    # remove all w:t in all runs, then put text in first run
    for r in runs:
        for t in r.findall(qn("w:t")):
            r.remove(t)
    for r in runs[1:]:
        p_el.remove(r)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    runs[0].append(t)


# ---------------------------------------------------------------------------
# Builder
# ---------------------------------------------------------------------------

class ReportBuilder:
    def __init__(self):
        self.doc = Document(str(TEMPLATE))
        self.fig_counters: dict[int, int] = {}
        self.tbl_counters: dict[int, int] = {}
        self.cur_chapter = 0
        self.figure_index: list[tuple[str, str]] = []   # (label, caption)
        self.table_index: list[tuple[str, str]] = []
        self._tune_styles()

    def _tune_styles(self):
        normal = self.doc.styles["Normal"]
        normal.font.name = BODY_FONT
        normal.font.size = Pt(BODY_SIZE)
        rpr = normal.element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        for attr in ("w:ascii", "w:hAnsi", "w:cs"):
            rfonts.set(qn(attr), BODY_FONT)
        normal.paragraph_format.line_spacing = 1.5
        normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

    # -- element factories (created at end, then relocated before an anchor) --

    def _new_paragraph(self):
        return self.doc.add_paragraph()

    def _relocate_before(self, anchor_p, element):
        anchor_p.addprevious(element)

    # -- public block inserters (anchor is a w:p element) --

    def insert_heading(self, anchor, text, level=2):
        # Apply the real Word heading style so the heading is part of the
        # document outline (navigation and table of contents), then set explicit
        # run formatting for a consistent look across the report.
        style = {2: "Heading 2", 3: "Heading 3"}.get(level, "Heading 2")
        p = self.doc.add_paragraph(style=style)
        run = p.add_run(text)
        sizes = {2: 14, 3: 12}
        set_run_font(run, size=sizes.get(level, 13), bold=True, color=(0x1F, 0x39, 0x64))
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        self._relocate_before(anchor, p._p)

    def insert_para(self, anchor, text, bold=False, italic=False):
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        set_run_font(run, bold=bold, italic=italic)
        style_paragraph(p)
        self._relocate_before(anchor, p._p)

    def insert_bullets(self, anchor, items, numbered=False):
        style = "List Number" if numbered else "List Bullet"
        for it in items:
            p = self.doc.add_paragraph(style=style)
            if isinstance(it, tuple):
                run = p.add_run(it[0] + ": ")
                set_run_font(run, bold=True)
                run2 = p.add_run(it[1])
                set_run_font(run2)
            else:
                run = p.add_run(it)
                set_run_font(run)
            p.paragraph_format.line_spacing = 1.3
            p.paragraph_format.space_after = Pt(3)
            self._relocate_before(anchor, p._p)

    def insert_figure(self, anchor, filename, caption, width=6.0):
        self.fig_counters[self.cur_chapter] = self.fig_counters.get(self.cur_chapter, 0) + 1
        label = f"Figure {self.cur_chapter}.{self.fig_counters[self.cur_chapter]}"
        path = FIGDIR / filename
        # Fit within the usable page area so nothing is cropped. With 1" margins
        # on A4 the printable region is roughly 6.3" wide x 9.2" tall; we cap to
        # 6.0" x 7.6" to leave room for the caption and surrounding text.
        MAX_W, MAX_H = 6.0, 7.6
        px_w, px_h = _png_size(path)
        ar = px_w / px_h if px_h else 1.0
        w = min(width, MAX_W)
        h = w / ar
        if h > MAX_H:                      # too tall -> constrain by height instead
            h = MAX_H
            w = h * ar
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(path), width=Inches(w), height=Inches(h))
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        self._relocate_before(anchor, p._p)

        cap = self.doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = cap.add_run(f"{label}: ")
        set_run_font(r1, size=10, bold=True)
        r2 = cap.add_run(caption)
        set_run_font(r2, size=10, italic=True)
        cap.paragraph_format.space_after = Pt(10)
        self._relocate_before(anchor, cap._p)
        self.figure_index.append((label, caption))

    def insert_table(self, anchor, headers, rows, caption, col_widths=None):
        self.tbl_counters[self.cur_chapter] = self.tbl_counters.get(self.cur_chapter, 0) + 1
        label = f"Table {self.cur_chapter}.{self.tbl_counters[self.cur_chapter]}"
        # caption above table
        cap = self.doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = cap.add_run(f"{label}: ")
        set_run_font(r1, size=10, bold=True)
        r2 = cap.add_run(caption)
        set_run_font(r2, size=10, italic=True)
        cap.paragraph_format.space_before = Pt(8)
        cap.paragraph_format.space_after = Pt(2)
        self._relocate_before(anchor, cap._p)

        table = self.doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr[i].text = ""
            run = hdr[i].paragraphs[0].add_run(h)
            set_run_font(run, size=10, bold=True)
            hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            _shade_cell(hdr[i], "D9E2F3")
        for row in rows:
            cells = table.add_row().cells
            for i, val in enumerate(row):
                cells[i].text = ""
                run = cells[i].paragraphs[0].add_run(str(val))
                set_run_font(run, size=10)
                cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._relocate_before(anchor, table._tbl)
        self.table_index.append((label, caption))

        spacer = self.doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(8)
        self._relocate_before(anchor, spacer._p)

    def insert_code(self, anchor, code, caption=None):
        if caption:
            cp = self.doc.add_paragraph()
            r = cp.add_run(caption)
            set_run_font(r, size=10, italic=True, bold=True)
            cp.paragraph_format.space_before = Pt(6)
            self._relocate_before(anchor, cp._p)
        p = self.doc.add_paragraph()
        run = p.add_run(code)
        set_run_font(run, name=MONO_FONT, size=9)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.space_after = Pt(8)
        _shade_paragraph(p, "F2F2F2")
        self._relocate_before(anchor, p._p)

    # -- render a list of content blocks before an anchor --

    def render_blocks(self, anchor, blocks):
        for b in blocks:
            kind = b[0]
            if kind == "h":
                self.insert_heading(anchor, b[1], b[2] if len(b) > 2 else 2)
            elif kind == "p":
                self.insert_para(anchor, b[1])
            elif kind == "b":
                self.insert_bullets(anchor, b[1], numbered=False)
            elif kind == "n":
                self.insert_bullets(anchor, b[1], numbered=True)
            elif kind == "fig":
                self.insert_figure(anchor, b[1], b[2], b[3] if len(b) > 3 else 6.0)
            elif kind == "tbl":
                self.insert_table(anchor, b[1], b[2], b[3])
            elif kind == "code":
                self.insert_code(anchor, b[1], b[2] if len(b) > 2 else None)


def _png_size(path):
    """Return (width_px, height_px) of a PNG without external dependencies."""
    import struct
    with open(path, "rb") as f:
        head = f.read(24)
    if head[:8] == b"\x89PNG\r\n\x1a\n":
        return struct.unpack(">II", head[16:24])
    return (1, 1)


def force_font(doc, font_name=BODY_FONT):
    """Set every run in the document (body, cover content control, tables) to
    one font, so the whole report is typographically uniform. The template mixes
    theme fonts (Cambria/Calibri) with Times New Roman; this unifies them."""
    for r in doc.element.body.iter(qn("w:r")):
        rpr = r.find(qn("w:rPr"))
        if rpr is None:
            rpr = OxmlElement("w:rPr")
            r.insert(0, rpr)
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.insert(0, rfonts)
        for attr in ("w:ascii", "w:hAnsi", "w:cs"):
            rfonts.set(qn(attr), font_name)


def unwrap_nested_sdts(doc):
    """Flatten nested content controls (w:sdt) in the cover. The title sits in a
    content control nested inside the cover control; some renderers display such
    a nested control's placeholder text ('Project Title ____') instead of the
    value we set. Unwrapping it — replacing the control with its content —
    turns the title into plain text that every renderer shows verbatim."""
    body = doc.element.body
    changed = True
    while changed:
        changed = False
        for sdt in body.findall(".//" + qn("w:sdt")):
            parent = sdt.getparent()
            anc = parent
            nested = False
            while anc is not None and anc is not body:
                if anc.tag == qn("w:sdt"):
                    nested = True
                    break
                anc = anc.getparent()
            if not nested:
                continue
            content = sdt.find(qn("w:sdtContent"))
            idx = list(parent).index(sdt)
            if content is not None:
                for child in list(content):
                    parent.insert(idx, child)
                    idx += 1
            parent.remove(sdt)
            changed = True
            break


def fix_page_breaks(doc):
    """Eliminate stray blank pages. The template starts each section with a
    standalone empty paragraph that carries a manual page break; when prior
    content happens to fill a page, that empty paragraph lands alone on a fresh
    page and produces a blank one. We remove those break paragraphs and instead
    set 'page break before' on each top-level heading, which starts a new page
    reliably without ever leaving a blank one. Trailing empty paragraphs are
    also removed so the document does not end on a blank page."""
    # 1. remove standalone page-break paragraphs
    for p in list(doc.paragraphs):
        if p.text.strip():
            continue
        brs = p._p.findall(".//" + qn("w:br"))
        if any(b.get(qn("w:type")) == "page" for b in brs):
            p._p.getparent().remove(p._p)
    # 2. page-break-before on each major heading (chapters + prelim sections),
    #    skipping the very first one (Certificate) because the cover already ends
    #    with its own break — adding another would create a blank page after it.
    first_h1_seen = False
    for p in doc.paragraphs:
        t = p.text.strip()
        if p.style.name == "Heading 1":
            if not first_h1_seen:
                first_h1_seen = True
                continue
            p.paragraph_format.page_break_before = True
        elif p.style.name == "Heading 2" and t.startswith("Declaration"):
            p.paragraph_format.page_break_before = True
    # 3. drop trailing empty paragraphs (but keep one holding the section props)
    paras = doc.paragraphs
    while len(paras) > 1 and not paras[-1].text.strip():
        last = paras[-1]._p
        if last.find(qn("w:pPr") + "/" + qn("w:sectPr")) is not None:
            break
        if last.findall(".//" + qn("w:drawing")):
            break
        last.getparent().remove(last)
        paras = doc.paragraphs


def _shade_cell(cell, hex_color):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tcpr.append(shd)


def _shade_paragraph(p, hex_color):
    ppr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    ppr.append(shd)


# ---------------------------------------------------------------------------
# Cover page + preliminary sections
# ---------------------------------------------------------------------------

def fill_cover_and_prelims(rb: ReportBuilder):
    doc = rb.doc
    body = doc.element.body
    sdt = body.findall(".//" + qn("w:sdt"))[0]
    cover_paras = sdt.findall(".//" + qn("w:p"))

    def set_cover(idx, text):
        set_element_text(cover_paras[idx], text)

    set_cover(1, C.PROJECT_TITLE)
    set_cover(4, "Master of Computer Applications (MCA)")
    set_cover(6, f"Student Name: {C.STUDENT_NAME}")
    set_cover(7, f"Enrollment No: {C.ENROLLMENT_NO}        Registration No: {C.REG_NO}")
    set_cover(9, f"Guide/Mentor Name: {C.GUIDE_NAME}")

    # Declaration + acknowledgement (body paragraphs, matched by text prefix)
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.startswith("I, ____"):
            set_element_text(p._p, C.DECLARATION_1)
        elif t.startswith("This project has been carried out"):
            set_element_text(p._p, C.DECLARATION_2)
        elif t.startswith("Place:"):
            set_element_text(p._p, f"Place: {C.PLACE}                    Date: {C.DATE}")
        elif t.startswith("Student Signature:"):
            set_element_text(p._p, C.DECLARATION_SIGN)
        elif t.startswith("Write a short paragraph thanking"):
            set_element_text(p._p, C.ACKNOWLEDGEMENT)
        for r in p.runs:
            pass


# ---------------------------------------------------------------------------
# Main build
# ---------------------------------------------------------------------------

def setup_page(doc):
    """Apply the formatting required by the writing guidelines: A4 paper, 1-inch
    margins on all sides, and a bottom-right page number."""
    for section in doc.sections:
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        # page numbers bottom-right
        for p in section.footer.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for jc in section.footer._element.iter(qn("w:jc")):
            jc.set(qn("w:val"), "right")


def normalize_template(rb: ReportBuilder):
    """Fix two defects in the supplied template: the Chapter 4 heading is not
    styled as a heading (so it would be missing from the auto-generated table of
    contents), and 'Implementation' is misspelt with an accent."""
    for p in rb.doc.paragraphs:
        t = p.text.strip()
        if t.startswith("Chapter 4: System Design") and p.style.name != "Heading 1":
            p.style = rb.doc.styles["Heading 1"]
        if "Implémentation" in t:
            set_element_text(p._p, t.replace("Implémentation", "Implementation"))
            p.style = rb.doc.styles["Heading 1"]


TOC_MARKER = "@@TOC_MARKER@@"
LOF_MARKER = "@@LOF_MARKER@@"
LOT_MARKER = "@@LOT_MARKER@@"

ABBREVIATIONS = [
    ("AI", "Artificial Intelligence"),
    ("AIOps", "Artificial Intelligence for IT Operations"),
    ("API", "Application Programming Interface"),
    ("AWS", "Amazon Web Services"),
    ("CPU", "Central Processing Unit"),
    ("CSV", "Comma-Separated Values"),
    ("DFD", "Data-Flow Diagram"),
    ("EC2", "Elastic Compute Cloud (AWS)"),
    ("ER", "Entity–Relationship"),
    ("F1", "F1-score (harmonic mean of precision and recall)"),
    ("FPR", "False-Positive Rate"),
    ("HTTP", "HyperText Transfer Protocol"),
    ("IF", "Isolation Forest"),
    ("JSON", "JavaScript Object Notation"),
    ("LSTM", "Long Short-Term Memory"),
    ("ML", "Machine Learning"),
    ("MSE", "Mean Squared Error"),
    ("NAB", "Numenta Anomaly Benchmark"),
    ("ORM", "Object–Relational Mapping"),
    ("PR", "Precision–Recall"),
    ("RDS", "Relational Database Service (AWS)"),
    ("REST", "Representational State Transfer"),
    ("RMS", "Root Mean Square"),
    ("SQL", "Structured Query Language"),
    ("TOC", "Table of Contents"),
    ("UML", "Unified Modeling Language"),
    ("YAML", "YAML Ain't Markup Language"),
]


def insert_toc(rb: ReportBuilder):
    """Replace the 'Use MS Word...' instruction with a uniquely identifiable
    marker paragraph. finalize_report.py later replaces the marker with a static
    Table of Contents whose page numbers are read from the rendered PDF."""
    for p in rb.doc.paragraphs:
        if p.text.strip().startswith("(Use MS Word"):
            set_element_text(p._p, TOC_MARKER)
            break


def _page_break_para(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)
    return p


def _heading1_para(doc, text):
    p = doc.add_paragraph(style="Heading 1")
    p.add_run(text)
    return p


def _marker_para(doc, marker):
    p = doc.add_paragraph()
    p.add_run(marker)
    return p


def insert_prelim_lists(rb: ReportBuilder):
    """Insert List of Figures, List of Tables and List of Abbreviations as
    preliminary pages after the Table of Contents (required by the guidelines).
    The figure/table lists are markers that finalize_report.py fills with
    captions and page numbers; the abbreviations list is built statically."""
    doc = rb.doc
    toc_marker = None
    for p in doc.paragraphs:
        if p.text.strip() == TOC_MARKER:
            toc_marker = p._p
            break
    if toc_marker is None:
        return

    ref = toc_marker  # insert each new element after this, advancing the cursor

    def after(el):
        ref_local[0].addnext(el)
        ref_local[0] = el

    ref_local = [ref]

    # --- List of Figures ---
    after(_page_break_para(doc)._p)
    after(_heading1_para(doc, "List of Figures")._p)
    after(_marker_para(doc, LOF_MARKER)._p)
    # --- List of Tables ---
    after(_page_break_para(doc)._p)
    after(_heading1_para(doc, "List of Tables")._p)
    after(_marker_para(doc, LOT_MARKER)._p)
    # --- List of Abbreviations (static table) ---
    after(_page_break_para(doc)._p)
    after(_heading1_para(doc, "List of Abbreviations")._p)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(("Abbreviation", "Expansion")):
        hdr[i].text = ""
        r = hdr[i].paragraphs[0].add_run(h)
        set_run_font(r, size=11, bold=True)
        _shade_cell(hdr[i], "D9E2F3")
    for abbr, exp in ABBREVIATIONS:
        cells = table.add_row().cells
        cells[0].text = ""; cells[1].text = ""
        r0 = cells[0].paragraphs[0].add_run(abbr); set_run_font(r0, size=11, bold=True)
        r1 = cells[1].paragraphs[0].add_run(exp); set_run_font(r1, size=11)
    after(table._tbl)


def main():
    rb = ReportBuilder()
    setup_page(rb.doc)
    normalize_template(rb)
    fill_cover_and_prelims(rb)
    unwrap_nested_sdts(rb.doc)   # flatten the title control so it renders as text

    # Collect the 11 "Write content here..." anchors in document order.
    anchors = [p for p in rb.doc.paragraphs if p.text.strip() == "Write content here..."]
    sections = C.SECTIONS  # list of (chapter_number, blocks) in order: abstract=0, ch1..ch10

    assert len(anchors) == len(sections), f"{len(anchors)} anchors vs {len(sections)} sections"

    for (chap_no, blocks), anchor in zip(sections, anchors):
        rb.cur_chapter = chap_no
        # The "(min pages...)" hint is the placeholder's previous sibling BEFORE
        # we insert anything; capture it now so we can delete it afterwards.
        hint = anchor._p.getprevious()
        rb.render_blocks(anchor._p, blocks)
        anchor._p.getparent().remove(anchor._p)
        if hint is not None:
            txt = "".join(t.text or "" for t in hint.findall(".//" + qn("w:t")))
            if txt.strip().startswith("("):
                hint.getparent().remove(hint)

    insert_toc(rb)
    insert_prelim_lists(rb)
    fix_page_breaks(rb.doc)   # remove blank pages; new page per major section
    force_font(rb.doc)        # unify typography across the whole document
    rb.doc.save(str(OUT))

    # Persist the figure/table indexes so finalize_report.py can build the
    # List of Figures and List of Tables with real page numbers.
    import json
    (ROOT / "report" / "_index.json").write_text(json.dumps({
        "figures": rb.figure_index,
        "tables": rb.table_index,
    }, ensure_ascii=False, indent=2))

    print(f"Saved {OUT.relative_to(ROOT)}")
    print(f"  figures embedded: {len(rb.figure_index)}, tables: {len(rb.table_index)}")


if __name__ == "__main__":
    main()
